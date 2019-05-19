import docx, re
from docx.enum.text import WD_COLOR_INDEX

# def highlight_text(filename, text):
#     document = docx.Document(filename)
#     for paragraph in document.paragraphs:
#         if text in paragraph.text:
#             for run in paragraph.runs:
#                 if text in run.text:
#                     x = run.text.split()
#                     run.clear()
#                     for i in x:
#                         run.add_text(i)
#                         if i in re.findall(REAL_NUMBER,text):
#                             run.font.highlight_color = WD_COLOR_INDEX.YELLOW

REAL_NUMBER = r'\d+\.\d+'  # 1.4
WEIRD_NUMBER1 = r'(\d+\.)'  # 42.
WEIRD_NUMBER2 = r'(\.\d+)'  # .42

def highlight_numbers(filename):
    doc = docx.Document(filename)
    for paragraph in doc.paragraphs:
        change_number(paragraph)
    doc.save('output.docx')

def change_number(paragraph):
    runs = list(map(lambda run: run.text, paragraph.runs))
    s = ''
    ordinary_numbers = re.findall(REAL_NUMBER,paragraph.text)
    #print(*ordinary_numbers, sep=' | ')
    #замена нефеншуйной фигни на правильную
    for i in range(len(runs)):
        for num in re.findall(WEIRD_NUMBER1, runs[i]):
            ind = runs[i].find(num)
            if ind < 0: continue
            nmb = runs[i][ind:ind+1+len(num)]
            f = True
            for n in ordinary_numbers:
                if nmb in n:
                    f = False
                    break
            if f:
                runs[i] = runs[i].replace(num, num + '0')
        for num in re.findall(WEIRD_NUMBER2, runs[i]):
            ind = runs[i].find(num)
            if ind < 0: continue
            elif ind == 0:
                runs[i] = runs[i].replace(num, '0' + num)
            else:
                nmb = runs[i][ind-1:ind+len(num)]
                f = True
                for n in ordinary_numbers:
                    if nmb in n:
                        f = False
                        break
                if num != '.0' and f:
                    runs[i] = runs[i].replace(num, '0' + num)
        s += runs[i]


    nums = re.findall(REAL_NUMBER, s)


    i = 0
    for num in nums:
        if i >= len(runs): break
        ind = runs[i].find(num)
        if ind < 0:
            i += 1
        elif ind == 0:
            run1 = num
            run2 = runs[i][ind+len(num):]
            del runs[i]
            runs.insert(i,run2)
            runs.insert(i,run1)
            i += 1
        else:
            run0 = runs[i][:ind]
            run1 = num
            run2 = runs[i][ind+len(num):]
            del runs[i]
            runs.insert(i,run2)
            runs.insert(i,run1)
            runs.insert(i,run0)
            i +=2

    # f = True
    # i = 0
    # j = 0
    # while f:
    #     num = nums[j]
    #     ind = runs[i].find(num)
    #     if ind < 0:
    #         j += 1
    #     elif ind == 0:
    #         run1 = num
    #         run2 = runs[i][ind:ind+len(num)]
    #         del runs[i]
    #         runs.insert(i,run2)
    #         runs.insert(i,run1)
    #         i += 1
    #     else:
    #         run0 = runs[i][:ind]
    #         run1 = num
    #         run2 = runs[i][ind:ind+len(num)]
    #         del runs[i]
    #         runs.insert(i,run2)
    #         runs.insert(i,run1)
    #         runs.insert(i,run0)
    #         i +=2
    #


    # print(*nums)
    # delim na chisla i drygoe

    # i = 0
    # l = len(runs)
    # while i<l:
    #     for num in nums:
    #         if num in runs[i]:
    #             KwLen = len(num)
    #             KwIndex = runs[i].find(num)
    #             run1 = runs[i][:KwIndex]
    #             run2 = runs[i][KwIndex: KwIndex + KwLen]
    #             run3 = runs[i][KwIndex + KwLen:]
    #             del runs[i]
    #             runs.insert(i, run3)
    #             runs.insert(i, run2)
    #             runs.insert(i, run1)
    #             l += 2
    #         i += 1
    ####a shto esli paragraph. add ne navorochena i vosprinimaet ili text ili run a ne bez raznici??
    #peredelali
    paragraph.clear()
    for run in runs:
        if run in nums:
            spec = paragraph.add_run(run)
            spec.font.highlight_color = WD_COLOR_INDEX.YELLOW
        else:
            paragraph.add_run(run)


    '''
    #make indexs for each run in my text
            lowPara = paragraph.text.lower()
            lowKeyword = keyword.lower()
            KwLen = len(lowKeyword)
            KwIndex = lowPara.find(lowKeyword)
            run1 = paragraph.text[:KwIndex]
            run2 = paragraph.text[KwIndex : KwIndex + KwLen]
            run3 = paragraph.text[KwIndex + KwLen:]

            #clear the current text and add in my highlighted keyword
            paragraph.text = ''
            paragraph.add_run(run1)
    '''

    # string = run.text.split()
    # run.clear()
    # for word in string:
    #     run.add_text(' '+ word)
    #     if word in nums:
    #         run.font.highlight_color = WD_COLOR_INDEX.YELLOW
    #
    # del string

# def change_number(run):
#     string = run.text
#     for num in re.findall(WEIRD_NUMBER1, string):
#         # print(num+'0')
#         string = string.replace(num, num + '0')
#     for num in re.findall(WEIRD_NUMBER2, string):
#         if num != '.0':
#             string = string.replace(num, '0' + num)
#
#     run.text = string
#     nums = re.findall(REAL_NUMBER, string)
#
#     string = run.text.split()
#     run.clear()
#     for word in string:
#         run.add_text(' '+ word)
#         if word in nums:
#             run.font.highlight_color = WD_COLOR_INDEX.YELLOW
#
#     del string


# def ChangeNumber(string):
#     for num in re.finditer(WEIRD_NUMBER1, string):
#         string.replace(num, num + '0')
#     for num in re.finditer(WEIRD_NUMBER2, string):
#         string.replace(num, '0' + num)

filename = 'input.docx'
highlight_numbers(filename)
